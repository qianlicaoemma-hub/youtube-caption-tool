from __future__ import annotations

import html
import json
import os
import re
import shutil
import subprocess
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Callable

from app import volcengine_asr

Progress = Callable[..., None]


@dataclass(frozen=True)
class JobOptions:
    url: str
    mode: str = "auto"
    language: str = "auto"
    cookies_from_browser: str | None = None
    cookies_file: str | None = None
    translate_to_zh: bool = False  # 保留字段，当前 UI 不暴露


@dataclass
class TranscriptBlock:
    speaker: str
    text: str


def run_transcription(
    options: JobOptions,
    work_dir: Path,
    output_dir: Path,
    progress: Progress,
) -> dict[str, Any]:
    work_dir.mkdir(parents=True, exist_ok=True)
    output_dir.mkdir(parents=True, exist_ok=True)

    completed = False
    try:
        meta = _get_metadata(options, progress)
        title = meta["title"]
        uploader = meta.get("uploader") or ""

        # ---- 字幕优先（auto / captions 模式） ----
        if options.mode in {"auto", "captions"}:
            try:
                progress("正在尝试读取 YouTube 自带字幕。")
                caption_blocks = _download_and_parse_captions(options, work_dir / "captions")
                if caption_blocks:
                    blocks = [
                        TranscriptBlock(block.speaker, _group_sentences(_light_cleanup(block.text)))
                        for block in caption_blocks
                    ]
                    files = _write_outputs(output_dir, title, options.url, blocks,
                                           "YouTube captions", uploader)
                    completed = True
                    return _result(title, options.url, "YouTube captions", blocks, files, uploader)
                if options.mode == "captions":
                    raise RuntimeError("这个视频没有可用的 YouTube 字幕。")
            except Exception as exc:
                if options.mode == "captions":
                    raise RuntimeError(f"读取 YouTube 字幕失败：{exc}") from exc
                progress("没有拿到可用字幕，切换到火山引擎音频识别。")

        # ---- 火山引擎音频识别 ----
        progress("正在准备用火山引擎/豆包识别音频。")
        blocks = _transcribe_audio_with_volcengine(options, work_dir, progress)
        files = _write_outputs(output_dir, title, options.url, blocks,
                               "火山引擎/豆包 ASR", uploader)
        completed = True
        return _result(title, options.url, "火山引擎/豆包 ASR", blocks, files, uploader)
    finally:
        if completed:
            shutil.rmtree(work_dir, ignore_errors=True)


def _result(
    title: str,
    url: str,
    source: str,
    blocks: list[TranscriptBlock],
    files: dict[str, str],
    uploader: str = "",
) -> dict[str, Any]:
    return {
        "title": title,
        "url": url,
        "uploader": uploader,
        "source": source,
        "text": _render_body(blocks),
        "files": files,
    }


# ---------- 火山引擎音频识别 ----------

def _transcribe_audio_with_volcengine(
    options: JobOptions,
    work_dir: Path,
    progress: Progress,
) -> list[TranscriptBlock]:
    audio_dir = work_dir / "audio"
    audio_dir.mkdir(parents=True, exist_ok=True)

    source_audio = _existing_source_audio(audio_dir)
    if source_audio:
        progress("发现上次保留的临时音频，直接复用。")
    else:
        progress("正在下载临时音频，不会下载视频画面。")
        source_audio = _download_audio(options, audio_dir)

    # 火山支持单文件最长 5 小时、512 MB，无需切片
    # 但需要压缩到合理体积（4 小时 16kHz mono 32kbps ≈ 60 MB）
    compressed_audio = audio_dir / "compressed.m4a"
    if not compressed_audio.exists():
        progress("正在压缩音频以降低上传体积。")
        _compress_audio(source_audio, compressed_audio)

    response = volcengine_asr.transcribe(
        audio_path=compressed_audio,
        work_dir=work_dir,
        progress=progress,
        enable_speakers=True,
    )
    pairs = volcengine_asr.parse_utterances(response)
    if not pairs:
        raise RuntimeError("火山引擎没有返回可识别的内容。")

    blocks = [TranscriptBlock(speaker, text) for speaker, text in pairs]
    merged = _merge_adjacent_blocks(blocks)
    return _clean_blocks(merged)


def _compress_audio(source_audio: Path, output_path: Path) -> None:
    """转成 16kHz mono 32kbps m4a，给火山上传用。"""
    command = [
        "ffmpeg",
        "-y",
        "-i",
        str(source_audio),
        "-vn",
        "-ac",
        "1",
        "-ar",
        "16000",
        "-c:a",
        "aac",
        "-b:a",
        "32k",
        str(output_path),
    ]
    _run(command, "压缩音频失败")


def _existing_source_audio(audio_dir: Path) -> Path | None:
    candidates = [path for path in audio_dir.glob("source.*") if path.is_file()]
    if not candidates:
        return None
    return max(candidates, key=lambda path: path.stat().st_size)


def _clean_blocks(blocks: list[TranscriptBlock]) -> list[TranscriptBlock]:
    return [TranscriptBlock(block.speaker, _group_sentences(_light_cleanup(block.text)))
            for block in blocks if block.text.strip()]


def _group_sentences(text: str, sentences_per_paragraph: int = 4) -> str:
    """把长文本按句号/问号/感叹号切句，每 N 句组成一段（用空行分隔）。

    - 中英文终止标点都识别（。！？.!?）
    - 标点保留在句尾
    - 短文本（< 2 句）原样返回，不强行分段
    """
    text = text.strip()
    if not text:
        return ""

    sentences = re.split(r"(?<=[。！？.!?])\s+", text)
    sentences = [s.strip() for s in sentences if s.strip()]

    if len(sentences) <= 1:
        return text

    paragraphs = []
    for i in range(0, len(sentences), sentences_per_paragraph):
        chunk = sentences[i:i + sentences_per_paragraph]
        paragraphs.append(" ".join(chunk))
    return "\n\n".join(paragraphs)


# ---------- yt-dlp：音频下载 + 字幕抓取 + 标题 ----------

def _download_audio(options: JobOptions, output_dir: Path) -> Path:
    output_template = output_dir / "source.%(ext)s"
    command = [
        "yt-dlp",
        "-f",
        "bestaudio/best",
        "-o",
        str(output_template),
        "--no-playlist",
        options.url,
    ]
    _run_ytdlp(command, options, "下载音频失败")

    candidates = [path for path in output_dir.glob("source.*") if path.is_file()]
    if not candidates:
        raise RuntimeError("没有找到下载后的音频文件。")
    return max(candidates, key=lambda path: path.stat().st_size)


def _download_and_parse_captions(options: JobOptions, output_dir: Path) -> list[TranscriptBlock]:
    shutil.rmtree(output_dir, ignore_errors=True)
    output_dir.mkdir(parents=True, exist_ok=True)

    output_template = output_dir / "%(title).120B-%(id)s.%(ext)s"
    command = [
        "yt-dlp",
        "--skip-download",
        "--write-subs",
        "--write-auto-subs",
        "--sub-langs",
        _subtitle_languages(options.language),
        "--sub-format",
        "vtt",
        "--no-playlist",
        "-o",
        str(output_template),
        options.url,
    ]
    _run_ytdlp(command, options, "读取字幕失败")

    vtt_files = sorted(output_dir.glob("*.vtt"), key=_caption_priority)
    if options.language == "auto":
        english_text = ""
        chinese_text = ""
        for path in vtt_files:
            text = _parse_vtt(path.read_text(encoding="utf-8", errors="ignore"))
            if len(text) <= 200:
                continue
            if not english_text and _is_english_caption(path):
                english_text = text
            elif not chinese_text and _is_chinese_caption(path):
                chinese_text = text
            if english_text and chinese_text:
                break

        blocks: list[TranscriptBlock] = []
        if english_text:
            blocks.append(TranscriptBlock("English", english_text))
        if chinese_text:
            blocks.append(TranscriptBlock("中文", chinese_text))
        if blocks:
            return blocks

    for path in vtt_files:
        text = _parse_vtt(path.read_text(encoding="utf-8", errors="ignore"))
        if len(text) > 200:
            return [TranscriptBlock("Speaker 1", text)]
    return []


def _subtitle_languages(language: str) -> str:
    if language == "zh":
        return "zh-CN,zh-TW,zh-Hans,zh-Hant,zh"
    if language == "en":
        return "en-orig,en"
    return "en-orig,en,zh-CN,zh-TW,zh-Hans,zh-Hant,zh"


def _caption_priority(path: Path) -> tuple[int, str]:
    name = path.name.lower()
    if ".en-orig" in name:
        return (0, name)
    if ".en" in name or "english" in name:
        return (1, name)
    if ".zh" in name or "chinese" in name:
        return (2, name)
    return (3, name)


def _is_english_caption(path: Path) -> bool:
    name = path.name.lower()
    return ".en-orig" in name or ".en." in name or name.endswith(".en.vtt") or "english" in name


def _is_chinese_caption(path: Path) -> bool:
    name = path.name.lower()
    return ".zh" in name or "chinese" in name


def _get_metadata(options: JobOptions, progress: Progress) -> dict[str, str]:
    """读取视频标题和上传者。"""
    command = ["yt-dlp", "--dump-json", "--skip-download", "--no-playlist", options.url]
    try:
        completed = _run_ytdlp(command, options, "读取视频信息失败", timeout=60)
        data = json.loads(completed.stdout)
        title = data.get("title") or "YouTube Transcript"
        uploader = data.get("uploader") or data.get("channel") or ""
        return {
            "title": str(title).strip() or "YouTube Transcript",
            "uploader": str(uploader).strip(),
        }
    except Exception:
        progress("暂时没有读到视频标题/作者信息，继续处理正文。")
        return {"title": "YouTube Transcript", "uploader": ""}


def _get_title(options: JobOptions, progress: Progress) -> str:
    """向后兼容的旧接口（保留以避免引入额外修改）。"""
    return _get_metadata(options, progress)["title"]


def _yt_dlp_attempts(command: list[str], options: JobOptions) -> list[list[str]]:
    base = [*command]
    if shutil.which("node"):
        base[1:1] = ["--js-runtimes", "node", "--remote-components", "ejs:github"]

    if _is_public_mode():
        return [base]

    if options.cookies_file:
        return [[*base, "--cookies", options.cookies_file]]

    if options.cookies_from_browser and options.cookies_from_browser != "auto":
        return [[*base, "--cookies-from-browser", options.cookies_from_browser]]

    attempts = [base]
    for browser in ("chrome:Profile 1", "chrome:Profile 12", "chrome:Default", "chrome:Profile 7", "safari", "firefox", "edge"):
        attempts.append([*base, "--cookies-from-browser", browser])
    return attempts


def _run_ytdlp(
    command: list[str],
    options: JobOptions,
    message: str,
    timeout: int | None = None,
) -> subprocess.CompletedProcess[str]:
    errors: list[str] = []
    for attempt in _yt_dlp_attempts(command, options):
        completed = subprocess.run(attempt, capture_output=True, text=True, timeout=timeout)
        if completed.returncode == 0:
            return completed
        detail = (completed.stderr or completed.stdout).strip()
        errors.append(detail[-1200:])

    final_error = errors[-1] if errors else "未知错误"
    if _is_public_mode():
        public_error = _public_youtube_error(final_error)
        if public_error:
            raise RuntimeError(f"{message}：{public_error}")
    raise RuntimeError(f"{message}：{final_error}")


def _is_public_mode() -> bool:
    return os.getenv("PUBLIC_MODE", "").strip().lower() in {"1", "true", "yes", "on"}


def _public_youtube_error(error: str) -> str:
    lowered = error.lower()
    if "429" in lowered or "too many requests" in lowered or "not a bot" in lowered:
        return (
            "YouTube 对当前公开服务器请求做了风控限制。"
            "请稍后重试，或换一个有公开字幕的视频。"
        )
    if "sign in" in lowered or "cookies" in lowered:
        return (
            "这个视频需要登录验证，公开版不会使用私人 YouTube 登录态。"
            "请换一个有公开字幕、无需登录验证的视频。"
        )
    return ""


def _run(command: list[str], message: str) -> subprocess.CompletedProcess[str]:
    completed = subprocess.run(command, capture_output=True, text=True)
    if completed.returncode != 0:
        detail = (completed.stderr or completed.stdout).strip()
        raise RuntimeError(f"{message}：{detail[-1200:]}")
    return completed


# ---------- VTT 字幕解析 ----------

def _parse_vtt(content: str) -> str:
    lines: list[str] = []
    last = ""
    skip_block = False

    for raw_line in content.splitlines():
        line = raw_line.strip()
        if not line:
            skip_block = False
            continue
        if line.startswith(("WEBVTT", "Kind:", "Language:")):
            continue
        if line.startswith("NOTE"):
            skip_block = True
            continue
        if skip_block:
            continue
        if "-->" in line:
            continue
        if line.isdigit():
            continue

        line = re.sub(r"<[^>]+>", "", line)
        line = html.unescape(line)
        line = re.sub(r"\s+", " ", line).strip()
        if not line or line == last:
            continue
        lines.append(line)
        last = line

    return _join_caption_lines(lines)


def _join_caption_lines(lines: list[str]) -> str:
    paragraphs: list[str] = []
    current: list[str] = []

    for line in lines:
        current.append(line)
        if re.search(r"[。！？.!?]$", line) or len(" ".join(current)) > 500:
            paragraphs.append(" ".join(current))
            current = []
    if current:
        paragraphs.append(" ".join(current))

    return "\n\n".join(paragraphs)


# ---------- 文本清洗与块合并 ----------

def _merge_adjacent_blocks(blocks: list[TranscriptBlock]) -> list[TranscriptBlock]:
    merged: list[TranscriptBlock] = []
    for block in blocks:
        text = block.text.strip()
        if not text:
            continue
        if merged and merged[-1].speaker == block.speaker:
            merged[-1].text = f"{merged[-1].text} {text}".strip()
        else:
            merged.append(TranscriptBlock(block.speaker, text))
    return merged


def _light_cleanup(text: str) -> str:
    """轻度清洗：去掉口水词的明显重复，不删有信息量的内容。"""
    cleaned = text.strip()
    # 移除背景音标记
    cleaned = re.sub(r"\[(music|applause|laughter|音乐|掌声|笑声)\]", "", cleaned, flags=re.IGNORECASE)
    # 英文填充词（只删孤立的，不删 "ah, I see" 这类有语气的）
    cleaned = re.sub(r"\b(um+|uh+|er+|hmm+)\b[,\s]*", "", cleaned, flags=re.IGNORECASE)
    # 中文填充词只删 3 次以上的连续重复（"嗯嗯嗯嗯"），保留单字
    cleaned = re.sub(r"(嗯|呃|啊|额){3,}", r"\1", cleaned)
    # 同词重复 3 次以上才压缩，保留 "very very good" 之类
    cleaned = re.sub(r"\b(\w+)(\s+\1\b){2,}", r"\1", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"([一-鿿])(?:\s*\1){2,}", r"\1", cleaned)
    cleaned = re.sub(r"[ \t]+", " ", cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


# ---------- 输出渲染 ----------

def _write_outputs(
    output_dir: Path,
    title: str,
    url: str,
    blocks: list[TranscriptBlock],
    source: str,
    uploader: str = "",
) -> dict[str, str]:
    output_dir.mkdir(parents=True, exist_ok=True)
    safe_name = _safe_filename(title)

    markdown = _render_markdown(title, url, source, blocks, uploader)
    text = _render_text(title, url, source, blocks, uploader)

    md_path = output_dir / f"{safe_name}.md"
    txt_path = output_dir / f"{safe_name}.txt"
    docx_path = output_dir / f"{safe_name}.docx"

    md_path.write_text(markdown, encoding="utf-8")
    txt_path.write_text(text, encoding="utf-8")
    _write_docx(docx_path, title, url, source, blocks, uploader)

    return {
        "markdown": str(md_path),
        "txt": str(txt_path),
        "docx": str(docx_path),
    }


def _meta_lines(url: str, source: str, uploader: str) -> str:
    lines = [f"Source: {url}"]
    if uploader:
        lines.append(f"Uploader: {uploader}")
    lines.append(f"Method: {source}")
    return "\n".join(lines)


def _render_markdown(title: str, url: str, source: str,
                     blocks: list[TranscriptBlock], uploader: str = "") -> str:
    return f"# {title}\n\n{_meta_lines(url, source, uploader)}\n\n{_render_body(blocks)}\n"


def _render_text(title: str, url: str, source: str,
                 blocks: list[TranscriptBlock], uploader: str = "") -> str:
    return f"{title}\n\n{_meta_lines(url, source, uploader)}\n\n{_render_body(blocks)}\n"


def _render_body(blocks: list[TranscriptBlock]) -> str:
    body = "\n\n".join(
        f"{block.speaker}:\n{block.text.strip()}"
        for block in blocks if block.text.strip()
    )
    return f"## Transcript\n\n{body}"


def _write_docx(path: Path, title: str, url: str, source: str,
                blocks: list[TranscriptBlock], uploader: str = "") -> None:
    from docx import Document

    document = Document()
    document.add_heading(title, level=1)
    document.add_paragraph(f"Source: {url}")
    if uploader:
        document.add_paragraph(f"Uploader: {uploader}")
    document.add_paragraph(f"Method: {source}")
    document.add_heading("Transcript", level=2)

    for block in blocks:
        speaker = document.add_paragraph()
        speaker.add_run(f"{block.speaker}:").bold = True
        # 把每段（按 \n\n 分隔）做成独立 paragraph，让 Word 里有视觉换行
        for paragraph_text in block.text.strip().split("\n\n"):
            paragraph_text = paragraph_text.strip()
            if paragraph_text:
                document.add_paragraph(paragraph_text)

    document.save(path)


def _safe_filename(name: str) -> str:
    cleaned = re.sub(r"[\\/:*?\"<>|]+", "-", name).strip()
    return cleaned[:100] or "youtube-transcript"
